"""DOCX generation with python-docx.

Renders the SAME deterministic context as the HTML templates, so PDF and DOCX
always agree on structure, numbering and grouping.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

HEADING_COLOR = RGBColor(0x1F, 0x29, 0x37)
MUTED = RGBColor(0x6B, 0x72, 0x80)


def _heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = HEADING_COLOR


def _table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(9)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            for p in cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9.5)
    doc.add_paragraph()


def context_to_docx(context: dict, out_path: Path) -> Path:
    """Build the DOCX from the template-engine context (already numbered/grouped)."""
    meeting = context["meeting"]
    doc = Document()

    # Header block
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(meeting.get("organization") or "Organization")
    run.font.size = Pt(16)
    run.font.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("MINUTES OF MEETING")
    run.font.size = Pt(11)
    run.font.color.rgb = MUTED

    _table(
        doc,
        ["Field", "Value", "Field", "Value"],
        [
            ["Meeting Title", meeting.get("title", ""), "Meeting Type", meeting.get("meeting_type", "")],
            ["Date", meeting.get("meeting_date") or "—", "Time", meeting.get("meeting_time") or "—"],
            [
                "Venue",
                meeting.get("venue") or "—",
                "Attendance",
                f"{context['attendee_stats']['present']} present / {context['attendee_stats']['absent']} absent",
            ],
        ],
    )

    if context.get("summary"):
        _heading(doc, "Summary")
        p = doc.add_paragraph(context["summary"])
        for run in p.runs:
            run.font.italic = True

    _heading(doc, "Attendees")
    for group in context["attendee_groups"]:
        _heading(doc, group["label"], level=2)
        _table(
            doc,
            ["#", "Name", "Role", "Department", "Status"],
            [
                [
                    str(i),
                    a["name"],
                    a["role"] or "—",
                    a["department"] or "—",
                    "Present" if a["present"] else "Absent",
                ]
                for i, a in enumerate(group["attendees"], start=1)
            ],
        )

    _heading(doc, "Agenda & Discussion")
    for item in context["agenda"]:
        _heading(doc, f"{item['number']}. {item['title']}", level=2)
        for sub in item["subtopics"]:
            doc.add_paragraph(f"{sub['number']}  {sub['title']}", style="List Bullet")
        for point in item["discussion"]:
            doc.add_paragraph(point, style="List Bullet 2")
    if context["unassigned_discussion"]:
        _heading(doc, "Other Discussion", level=2)
        for point in context["unassigned_discussion"]:
            doc.add_paragraph(point, style="List Bullet")

    if context["decisions"]:
        _heading(doc, "Decisions")
        _table(
            doc,
            ["#", "Decision", "Decided By", "Rationale"],
            [[d["number"], d["description"], d["decided_by"], d["rationale"]] for d in context["decisions"]],
        )

    if context["action_items"]:
        _heading(doc, "Action Items")
        _table(
            doc,
            ["#", "Action", "Owner", "Due Date", "Priority", "Status"],
            [
                [a["number"], a["description"], a["owner"], a["due_date"], a["priority"], a["status"]]
                for a in context["action_items"]
            ],
        )

    doc.add_paragraph()
    sig = doc.add_table(rows=2, cols=2)
    sig.alignment = WD_TABLE_ALIGNMENT.CENTER
    sig.rows[0].cells[0].text = meeting.get("prepared_by") or "—"
    sig.rows[0].cells[1].text = meeting.get("approved_by") or "—"
    sig.rows[1].cells[0].text = "Prepared By"
    sig.rows[1].cells[1].text = "Approved By"

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path
